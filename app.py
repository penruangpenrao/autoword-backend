from flask import Flask, request, send_file, render_template, jsonify
from flask_cors import CORS
import docx
from docx.shared import RGBColor
import io
import json

app = Flask(__name__)
CORS(app)

# 1. ฟังก์ชันเช็คสีแดง
def is_reddish(run):
    if run.font.color and run.font.color.rgb:
        hex_color = str(run.font.color.rgb).upper()
        try:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            if r > 130 and g < 100 and b < 100:
                return True
        except:
            pass
    return False

# 2. ฟังก์ชันกาววิเศษ (รวมคำสีแดงที่อยู่ติดกันให้เป็นประโยคเดียว)
def normalize_red_runs(doc):
    def _normalize_paragraph(p):
        first_red_run = None
        for run in p.runs:
            if is_reddish(run):
                if first_red_run is None:
                    first_red_run = run # เจอสีแดงก้อนแรก ให้จำไว้
                else:
                    # เจอสีแดงก้อนถัดมา เอาข้อความมาต่อท้ายก้อนแรก แล้วลบก้อนนี้ทิ้ง
                    first_red_run.text += run.text
                    run.text = ""
            else:
                first_red_run = None # ถ้าเจอสีดำ ให้ตัดจบการรวมร่าง

    # วนลูปติดกาวให้ทุกย่อหน้าและทุกตาราง
    for p in doc.paragraphs:
        _normalize_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _normalize_paragraph(p)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    import traceback
    from docx.oxml.ns import qn

    file = request.files.get('file')
    if not file:
        return jsonify({'error': 'ไม่พบไฟล์'}), 400

    try:
        doc = docx.Document(file)
        normalize_red_runs(doc)

        W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        nsmap  = {'w': W_NS}
        W_BODY = qn('w:body')

        # ── Step 1: คำนวณหน้าของแต่ละ body-child element ──
        # ใช้เฉพาะ hard page break (w:br w:type="page") เท่านั้น — เสถียรกว่า
        current_page = 1
        body_el_page = {}

        for body_child in doc.element.body:
            body_el_page[id(body_child)] = current_page
            pbreaks = body_child.xpath(
                './/w:br[@w:type="page"]',
                namespaces=nsmap
            )
            current_page += len(pbreaks)

        # ── Step 2: หาหน้าของ paragraph โดย walk-up ──
        def page_of(para):
            el = para._element
            while el is not None:
                parent = el.getparent()
                if parent is not None and parent.tag == W_BODY:
                    return body_el_page.get(id(el), 1)
                el = parent
            return 1

        # ── Step 3: ดึงคำสีแดง + หน้า ──
        word_pages = {}
        seen_order = []

        def extract(paragraphs):
            for p in paragraphs:
                pg = page_of(p)
                for run in p.runs:
                    if is_reddish(run):
                        text = run.text.strip()
                        if text:
                            if text not in word_pages:
                                word_pages[text] = set()
                                seen_order.append(text)
                            word_pages[text].add(pg)

        extract(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    extract(cell.paragraphs)

        result = [
            {'word': w, 'pages': sorted(word_pages[w])}
            for w in seen_order
        ]
        return jsonify({'words': result})

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': 'เกิดข้อผิดพลาด: ' + str(e)}), 500

@app.route('/generate', methods=['POST'])
def generate():
    file = request.files.get('file')
    replacements_json = request.form.get('replacements')
    
    if not file or not replacements_json:
        return jsonify({'error': 'ข้อมูลไม่ครบถ้วน'}), 400
        
    replacements = json.loads(replacements_json)
    doc = docx.Document(file)
    
    normalize_red_runs(doc) # <--- สั่งรวมคำก่อนทำการแก้คำด้วย

    def replace_and_recolor(paragraphs, replacements):
        for p in paragraphs:
            for run in p.runs:
                if is_reddish(run):
                    original = run.text.strip()
                    if original in replacements:
                        new_text = replacements[original]
                        # ถ้าช่องกรอกข้อมูลไม่ว่างเปล่า ให้แทนที่คำและเปลี่ยนเป็นสีดำ
                        if new_text != "": 
                            run.text = run.text.replace(original, new_text)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        # ถ้าเป็นค่าว่าง โค้ดจะข้ามไปและปล่อยให้เป็นสีแดงตามเดิม
                        
    replace_and_recolor(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_and_recolor(cell.paragraphs, replacements)
                                    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0) 
    
    return send_file(
        bio,
        as_attachment=True,
        download_name='เอกสาร_อัตโนมัติ.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=True, port=5000)
